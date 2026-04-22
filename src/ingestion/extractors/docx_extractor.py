"""
docx_extractor.py — Extractor for .docx files.

Extracts text from paragraphs and table cells using python-docx.
Both content types are included — tables often contain key information.
"""

from pathlib import Path
import docx


def extract(file_path: Path) -> str:
    """
    Extract plain text from a Word document.

    Args:
        file_path: Path to the .docx file.

    Returns:
        Extracted text as a single string. Paragraphs and table rows
        are separated by newlines.

    Raises:
        ValueError: If the file extension is not .docx.
    """
    if file_path.suffix.lower() != ".docx":
        raise ValueError(f"docx_extractor only handles .docx, got: {file_path.suffix}")

    document = docx.Document(file_path)
    blocks = []

    # extract main body paragraphs
    for para in document.paragraphs:
        text = para.text.strip()
        if text:                               # skip empty paragraphs
            blocks.append(text)

    # extract table cells — tables sit alongside paragraphs, not inside them
    for table in document.tables:
        for row in table.rows:
            # join cells in a row with pipe separator
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()           # skip empty cells
            )
            if row_text:
                blocks.append(row_text)

    return "\n".join(blocks)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python -m src.ingestion.extractors.docx_extractor <file.docx>")
        sys.exit(1)

    target = Path(sys.argv[1])

    if not target.exists():
        print(f"Error: file not found — {target}")
        sys.exit(1)

    result = extract(target)

    print(f"--- Extracted {len(result):,} characters from {target.name} ---\n")
    print(result[:500])